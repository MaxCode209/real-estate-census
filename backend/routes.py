"""API routes for the application."""
from flask import Blueprint, jsonify, request
from sqlalchemy.orm import Session, load_only
from sqlalchemy import or_, text
from typing import List, Dict, Optional
from backend.database import get_db
from backend.models import CensusData, SchoolData, AttendanceZone

# Columns that exist in census_data table (no city until added in Supabase)
_CENSUS_LOAD_COLUMNS = (
    CensusData.id, CensusData.zip_code, CensusData.state, CensusData.county,
    CensusData.population, CensusData.median_age, CensusData.average_household_income,
    CensusData.data_year, CensusData.created_at, CensusData.updated_at,
)
# Explicit column list for raw SQL (never includes city)
_CENSUS_SQL_COLS = "id, zip_code, state, county, population, median_age, average_household_income, data_year, created_at, updated_at"
from backend.census_api import CensusAPIClient
from backend.zone_utils import (
    point_in_polygon,
    find_zoned_schools,
    find_all_zoned_schools,
    load_zip_polygon,
    zones_intersecting_zip,
    zones_intersecting_zip_diagnostic,
    group_zones_by_district,
    district_geometry_in_zip,
)
from backend.greatschools_client import GreatSchoolsClient

api = Blueprint('api', __name__, url_prefix='/api')

def _census_kwargs(data):
    """Filter dict to only keys that exist on CensusData model (avoids extra columns from API)."""
    allowed = {c.key for c in CensusData.__table__.c}
    return {k: v for k, v in (data or {}).items() if k in allowed}

@api.route('/census-data', methods=['GET'])
def get_census_data():
    """Get census data with optional filters. Uses raw SQL so we never reference city column."""
    try:
        db: Session = next(get_db())
    except Exception as e:
        return jsonify({'error': f'Database connection failed: {str(e)}', 'data': []}), 500

    try:
        zip_code = request.args.get('zip_code')
        min_income = request.args.get('min_income', type=float)
        max_income = request.args.get('max_income', type=float)
        min_population = request.args.get('min_population', type=int)
        max_population = request.args.get('max_population', type=int)
        limit = request.args.get('limit', type=int, default=1000)
        offset = request.args.get('offset', type=int, default=0)

        # Build WHERE and params (no ORM - avoids census_data.city)
        where_parts = []
        params = {}
        if zip_code:
            where_parts.append("zip_code = :zip_code")
            params["zip_code"] = zip_code
        if min_income:
            where_parts.append("average_household_income >= :min_income")
            params["min_income"] = min_income
        if max_income:
            where_parts.append("average_household_income <= :max_income")
            params["max_income"] = max_income
        if min_population:
            where_parts.append("population >= :min_population")
            params["min_population"] = min_population
        if max_population:
            where_parts.append("population <= :max_population")
            params["max_population"] = max_population
        where_sql = " AND ".join(where_parts) if where_parts else "1=1"

        # Count with raw SQL (never selects city)
        count_sql = text(f"SELECT COUNT(*) FROM census_data WHERE {where_sql}")
        total = db.execute(count_sql, params).scalar()

        # Data with raw SQL (only columns that exist)
        data_sql = text(
            f"SELECT {_CENSUS_SQL_COLS} FROM census_data WHERE {where_sql} "
            "ORDER BY zip_code LIMIT :lim OFFSET :off"
        )
        params["lim"] = limit
        params["off"] = offset
        rows = db.execute(data_sql, params).fetchall()

        # Build response dicts (same shape as to_dict)
        keys = ["id", "zip_code", "state", "county", "population", "median_age", "average_household_income", "data_year", "created_at", "updated_at"]
        data = [dict(zip(keys, row)) for row in rows]
        for row in data:
            if row.get("created_at"):
                row["created_at"] = row["created_at"].isoformat() if hasattr(row["created_at"], "isoformat") else str(row["created_at"])
            if row.get("updated_at"):
                row["updated_at"] = row["updated_at"].isoformat() if hasattr(row["updated_at"], "isoformat") else str(row["updated_at"]) if row["updated_at"] else None

        return jsonify({
            "data": data,
            "total": total,
            "limit": limit,
            "offset": offset,
        })
    except Exception as e:
        return jsonify({"error": str(e), "data": []}), 500

@api.route('/census-data/zip/<zip_code>', methods=['GET'])
def get_census_data_by_zip(zip_code: str):
    """Get census data for a specific zip code. Fetches from Census API if not in database."""
    db: Session = next(get_db())
    
    record = db.query(CensusData).options(load_only(*_CENSUS_LOAD_COLUMNS)).filter(CensusData.zip_code == zip_code).first()
    
    if not record:
        # Try to fetch from Census API
        print(f"[INFO] Zip code {zip_code} not in database, fetching from Census API...")
        try:
            client = CensusAPIClient()
            census_data_list = client.fetch_zip_code_data([zip_code])
            
            if census_data_list and len(census_data_list) > 0:
                data = census_data_list[0]
                # Store in database (only columns that exist on model)
                new_record = CensusData(**_census_kwargs(data))
                db.add(new_record)
                db.commit()
                print(f"[INFO] Successfully fetched and stored census data for zip {zip_code}")
                return jsonify(new_record.to_dict())
            else:
                print(f"[WARN] Census API returned no data for zip {zip_code}")
                return jsonify({'error': 'Zip code not found in Census API'}), 404
        except Exception as e:
            print(f"[ERROR] Failed to fetch census data for zip {zip_code}: {e}")
            return jsonify({'error': f'Failed to fetch census data: {str(e)}'}), 500
    
    return jsonify(record.to_dict())

@api.route('/geocode-zip/<zip_code>', methods=['GET'])
def geocode_zip(zip_code: str):
    """Backend geocoding endpoint for zip codes."""
    try:
        import requests
        from config.config import Config
        
        # Use Google Geocoding API via backend
        # This helps if frontend API key has restrictions
        geocode_url = 'https://maps.googleapis.com/maps/api/geocode/json'
        params = {
            'address': zip_code,
            'key': Config.GOOGLE_MAPS_API_KEY,
            'components': 'country:US'
        }
        
        response = requests.get(geocode_url, params=params, timeout=10)
        data = response.json()
        
        if data['status'] == 'OK' and data['results']:
            result = data['results'][0]
            geometry = result['geometry']
            
            return jsonify({
                'success': True,
                'location': {
                    'lat': geometry['location']['lat'],
                    'lng': geometry['location']['lng']
                },
                'bounds': geometry.get('bounds'),
                'viewport': geometry.get('viewport')
            })
        else:
            return jsonify({
                'success': False,
                'error': data.get('error_message', data.get('status', 'Unknown error'))
            }), 400
            
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@api.route('/zip-boundary/<zip_code>', methods=['GET'])
def get_zip_boundary(zip_code: str):
    """Get GeoJSON boundary polygon for a zip code."""
    try:
        import requests
        import json
        from pathlib import Path
        
        # FIRST: Check if we have a locally stored boundary (fastest, most reliable)
        try:
            boundary_file = Path('data/zip_boundaries') / f"{zip_code}.geojson"
            if boundary_file.exists():
                with open(boundary_file, 'r') as f:
                    data = json.load(f)
                    return jsonify(data)
        except Exception as e:
            print(f"Local boundary file check failed: {e}")
        
        # Try multiple sources for zip code boundaries
        # Source 1: OpenDataSoft API (most reliable)
        try:
            ods_url = "https://public.opendatasoft.com/api/explore/v2.1/catalog/datasets/us-zip-code-labels-and-boundaries/records"
            # Try with ZCTA5CE10 field (Census format)
            params = {
                'where': f'zcta5ce10="{zip_code}"',
                'limit': 1,
                'select': 'zcta5ce10,geo_shape'
            }
            response = requests.get(ods_url, params=params, timeout=15)
            if response.status_code == 200:
                data = response.json()
                if 'results' in data and len(data['results']) > 0:
                    record = data['results'][0]
                    if 'record' in record and 'fields' in record['record']:
                        fields = record['record']['fields']
                        if 'geo_shape' in fields:
                            geometry = fields['geo_shape']
                            geojson = {
                                'type': 'FeatureCollection',
                                'features': [{
                                    'type': 'Feature',
                                    'geometry': geometry,
                                    'properties': {'ZCTA5CE10': zip_code}
                                }]
                            }
                            try:
                                boundary_dir = Path('data/zip_boundaries')
                                boundary_dir.mkdir(parents=True, exist_ok=True)
                                boundary_file = boundary_dir / f"{zip_code}.geojson"
                                with open(boundary_file, 'w') as f:
                                    json.dump(geojson, f)
                            except Exception:
                                pass
                            return jsonify(geojson)
            
            # Try with zip_code field
            params2 = {
                'where': f'zip_code="{zip_code}"',
                'limit': 1,
                'select': 'zip_code,geo_shape'
            }
            response2 = requests.get(ods_url, params=params2, timeout=15)
            if response2.status_code == 200:
                data2 = response2.json()
                if 'results' in data2 and len(data2['results']) > 0:
                    record = data2['results'][0]
                    if 'record' in record and 'fields' in record['record']:
                        fields = record['record']['fields']
                        if 'geo_shape' in fields:
                            geometry = fields['geo_shape']
                            geojson = {
                                'type': 'FeatureCollection',
                                'features': [{
                                    'type': 'Feature',
                                    'geometry': geometry,
                                    'properties': {'ZCTA5CE10': zip_code}
                                }]
                            }
                            try:
                                boundary_dir = Path('data/zip_boundaries')
                                boundary_dir.mkdir(parents=True, exist_ok=True)
                                boundary_file = boundary_dir / f"{zip_code}.geojson"
                                with open(boundary_file, 'w') as f:
                                    json.dump(geojson, f)
                            except Exception:
                                pass
                            return jsonify(geojson)
        except Exception as e:
            print(f"OpenDataSoft failed: {e}")
        
        # Source 2: Try boundaries.io API (FREE tier available)
        try:
            from config.config import Config
            boundaries_api_key = getattr(Config, 'BOUNDARIES_IO_API_KEY', None)
            
            if boundaries_api_key:
                # Use Boundaries.io API (requires free API key)
                boundaries_url = f"https://boundaries.io/api/v1/boundary"
                params = {
                    'zipcode': zip_code,
                    'api_key': boundaries_api_key,
                    'format': 'geojson'
                }
                response = requests.get(boundaries_url, params=params, timeout=15)
                if response.status_code == 200:
                    data = response.json()
                    # Save locally for future use
                    try:
                        boundary_dir = Path('data/zip_boundaries')
                        boundary_dir.mkdir(parents=True, exist_ok=True)
                        boundary_file = boundary_dir / f"{zip_code}.geojson"
                        with open(boundary_file, 'w') as f:
                            json.dump(data, f)
                    except:
                        pass
                    return jsonify(data)
            else:
                # Try free endpoint (may have rate limits)
                boundaries_url = f"https://boundaries-io.herokuapp.com/zip/{zip_code}"
                response = requests.get(boundaries_url, timeout=15)
                if response.status_code == 200:
                    data = response.json()
                    if 'boundaries' in data or 'geometry' in data:
                        return jsonify(data)
        except Exception as e:
            print(f"Boundaries.io failed: {e}")
        
        # Source 3: Try GitHub repository (state-based)
        try:
            # Try different GitHub sources
            github_sources = [
                f"https://raw.githubusercontent.com/OpenDataDE/State-zip-code-GeoJSON/master/{zip_code[0]}/{zip_code}_polygon.geojson",
                f"https://raw.githubusercontent.com/OpenDataDE/State-zip-code-GeoJSON/master/zcta5/{zip_code}_polygon.geojson",
            ]
            
            for github_url in github_sources:
                response = requests.get(github_url, timeout=10)
                if response.status_code == 200:
                    data = response.json()
                    if 'type' in data and (data['type'] == 'FeatureCollection' or data['type'] == 'Feature'):
                        if data['type'] == 'Feature':
                            # Wrap single feature in FeatureCollection
                            return jsonify({
                                'type': 'FeatureCollection',
                                'features': [data]
                            })
                        return jsonify(data)
        except Exception as e:
            print(f"GitHub source failed: {e}")
        
        # Source 4: Try using Census TIGERweb (ZCTA layer is 2, not 82!)
        try:
            feature_url = "https://tigerweb.geo.census.gov/arcgis/rest/services/TIGERweb/tigerWMS_Current/MapServer/2/query"
            
            # Use objectIds approach - first find the object ID
            # Then query by object ID
            params_list = [
                {'where': f"ZCTA5='{zip_code}'", 'outFields': '*', 'f': 'geojson', 'outSR': '4326', 'returnGeometry': 'true'},
                {'where': f"ZCTA5 = '{zip_code}'", 'outFields': '*', 'f': 'geojson', 'outSR': '4326', 'returnGeometry': 'true'},
                {'where': f"ZCTA5CE10='{zip_code}'", 'outFields': '*', 'f': 'geojson', 'outSR': '4326', 'returnGeometry': 'true'},
            ]
            
            for params in params_list:
                try:
                    response = requests.get(feature_url, params=params, timeout=25)
                    if response.status_code == 200:
                        data = response.json()
                        # Check for actual features, not errors
                        if 'features' in data and isinstance(data['features'], list) and len(data['features']) > 0:
                            # Validate it's actually a polygon
                            feature = data['features'][0]
                            if 'geometry' in feature and feature['geometry'].get('type') in ['Polygon', 'MultiPolygon']:
                                # Save locally for future use
                                try:
                                    boundary_dir = Path('data/zip_boundaries')
                                    boundary_dir.mkdir(parents=True, exist_ok=True)
                                    boundary_file = boundary_dir / f"{zip_code}.geojson"
                                    with open(boundary_file, 'w') as f:
                                        json.dump(data, f)
                                except:
                                    pass
                                return jsonify(data)
                except:
                    continue
        except Exception as e:
            print(f"Census TIGER failed: {e}")
        
        # Source 5: Try alternative Census TIGERweb endpoints (all FREE)
        census_endpoints = [
            "https://tigerweb.geo.census.gov/arcgis/rest/services/TIGERweb/tigerWMS_ACS2022/MapServer/2/query",
            "https://tigerweb.geo.census.gov/arcgis/rest/services/TIGERweb/tigerWMS_ACS2021/MapServer/2/query",
        ]
        
        for alt_url in census_endpoints:
            try:
                params = {
                    'where': f"ZCTA5='{zip_code}'",
                    'outFields': '*',
                    'f': 'geojson',
                    'outSR': '4326',
                    'returnGeometry': 'true'
                }
                response = requests.get(alt_url, params=params, timeout=20)
                if response.status_code == 200:
                    data = response.json()
                    if 'features' in data and len(data['features']) > 0:
                        try:
                            boundary_dir = Path('data/zip_boundaries')
                            boundary_dir.mkdir(parents=True, exist_ok=True)
                            boundary_file = boundary_dir / f"{zip_code}.geojson"
                            with open(boundary_file, 'w') as f:
                                json.dump(data, f)
                        except Exception:
                            pass
                        return jsonify(data)
            except Exception as e:
                print(f"Census endpoint {alt_url} failed: {e}")
                continue
        
        # If all sources fail, return 404 (frontend will use approximate boundary)
        return jsonify({
            'error': 'Boundary not found',
            'message': f'Could not fetch exact boundary for zip code {zip_code}. Using approximate boundary from geocoding.'
        }), 404
        
    except Exception as e:
        print(f"Error in get_zip_boundary: {e}")
        return jsonify({
            'error': str(e)
        }), 500

@api.route('/census-data', methods=['POST'])
def add_census_data():
    """Add or update census data."""
    db: Session = next(get_db())
    
    data = request.get_json()
    
    if not data or 'zip_code' not in data:
        return jsonify({'error': 'zip_code is required'}), 400
    
    # Check if record exists (load_only so we never SELECT city - column may not exist in DB yet)
    existing = db.query(CensusData).options(load_only(*_CENSUS_LOAD_COLUMNS)).filter(
        CensusData.zip_code == data['zip_code']
    ).first()
    
    if existing:
        # Update existing record
        for key, value in data.items():
            if hasattr(existing, key) and key != 'id':
                setattr(existing, key, value)
    else:
        # Create new record (only columns that exist on model)
        existing = CensusData(**_census_kwargs(data))
        db.add(existing)
    
    db.commit()
    db.refresh(existing)
    
    return jsonify(existing.to_dict()), 201 if not existing.id else 200

@api.route('/census-data/bulk', methods=['POST'])
def add_census_data_bulk():
    """Add multiple census records at once."""
    db: Session = next(get_db())
    
    data_list = request.get_json()
    
    if not isinstance(data_list, list):
        return jsonify({'error': 'Expected a list of records'}), 400
    
    added = 0
    updated = 0
    
    for data in data_list:
        if 'zip_code' not in data:
            continue
        
        existing = db.query(CensusData).options(load_only(*_CENSUS_LOAD_COLUMNS)).filter(
            CensusData.zip_code == data['zip_code']
        ).first()
        
        if existing:
            for key, value in data.items():
                if hasattr(existing, key) and key != 'id':
                    setattr(existing, key, value)
            updated += 1
        else:
            new_record = CensusData(**_census_kwargs(data))
            db.add(new_record)
            added += 1
    
    db.commit()
    
    return jsonify({
        'message': 'Bulk update completed',
        'added': added,
        'updated': updated
    })

@api.route('/census-data/fetch', methods=['POST'])
def fetch_census_data():
    """Fetch census data from Census Bureau API and store in database."""
    db: Session = next(get_db())
    
    request_data = request.get_json() or {}
    zip_codes = request_data.get('zip_codes')  # Optional list of zip codes
    
    client = CensusAPIClient()
    census_data = client.fetch_zip_code_data(zip_codes)
    
    if not census_data:
        return jsonify({'error': 'No data fetched from Census API'}), 400
    
    # Store in database
    added = 0
    updated = 0
    
    for data in census_data:
        existing = db.query(CensusData).options(load_only(*_CENSUS_LOAD_COLUMNS)).filter(
            CensusData.zip_code == data['zip_code']
        ).first()
        
        if existing:
            for key, value in data.items():
                if hasattr(existing, key):
                    setattr(existing, key, value)
            updated += 1
        else:
            new_record = CensusData(**_census_kwargs(data))
            db.add(new_record)
            added += 1
    
    db.commit()
    
    return jsonify({
        'message': 'Census data fetched and stored',
        'added': added,
        'updated': updated,
        'total_fetched': len(census_data)
    })

@api.route('/export/csv', methods=['GET'])
def export_to_csv():
    """Export census data to CSV file (download)."""
    from flask import Response
    import csv
    import io
    from datetime import datetime
    
    try:
        db: Session = next(get_db())
        
        # Get filter parameters from request (same as /api/census-data endpoint)
        zip_code = request.args.get('zip_code')
        min_income = request.args.get('min_income', type=float)
        max_income = request.args.get('max_income', type=float)
        min_population = request.args.get('min_population', type=int)
        max_population = request.args.get('max_population', type=int)
        limit = request.args.get('limit', type=int, default=10000)
        
        # Build query with filters (load_only so we never SELECT city - column may not exist in DB yet)
        query = db.query(CensusData).options(load_only(*_CENSUS_LOAD_COLUMNS))
        
        if zip_code:
            query = query.filter(CensusData.zip_code == zip_code)
        if min_income:
            query = query.filter(CensusData.average_household_income >= min_income)
        if max_income:
            query = query.filter(CensusData.average_household_income <= max_income)
        if min_population:
            query = query.filter(CensusData.population >= min_population)
        if max_population:
            query = query.filter(CensusData.population <= max_population)
        
        # Apply limit and get records
        records = query.limit(limit).all()
        
        # Convert to list of dicts
        data = [record.to_dict() for record in records]
        
        if not data:
            # Provide more helpful error message
            if zip_code:
                return jsonify({
                    'error': 'No data to export',
                    'message': f'Zip code {zip_code} not found in database. This zip code may not have census data available.',
                    'zip_code': zip_code
                }), 400
            else:
                return jsonify({
                    'error': 'No data to export',
                    'message': 'No records match the specified filters.'
                }), 400
        
        # Create CSV in memory
        output = io.StringIO()
        
        # Get headers from first record
        headers = list(data[0].keys())
        writer = csv.DictWriter(output, fieldnames=headers)
        
        # Write headers
        writer.writeheader()
        
        # Write data rows
        for record in data:
            writer.writerow(record)
        
        # Create filename with timestamp and zip code if applicable
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        if zip_code:
            filename = f'census_data_{zip_code}_{timestamp}.csv'
        else:
            filename = f'census_data_{timestamp}.csv'
        
        # Create response with CSV data
        response = Response(
            output.getvalue(),
            mimetype='text/csv',
            headers={
                'Content-Disposition': f'attachment; filename="{filename}"'
            }
        )
        
        return response
        
    except ImportError as ie:
        missing_lib = str(ie)
        if 'googleapiclient' in missing_lib or 'google-api-python-client' in missing_lib.lower():
            return jsonify({
                'error': 'Google API client library not installed',
                'message': 'Install with: pip install google-api-python-client'
            }), 500
        else:
            return jsonify({
                'error': 'Google Sheets libraries not installed',
                'message': 'Install with: pip install gspread google-auth google-auth-oauthlib google-api-python-client'
            }), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@api.route('/export/report', methods=['GET'])
def export_report():
    """Export site report as Word document or PDF with demographics and top 10 schools."""
    try:
        from flask import Response
        from datetime import datetime
        from io import BytesIO
        import requests
        from config.config import Config
        from sqlalchemy import text
        
        # Get parameters
        address = request.args.get('address')
        lat = request.args.get('lat', type=float)
        lng = request.args.get('lng', type=float)
        zip_code = request.args.get('zip_code')
        format_type = request.args.get('format', 'docx')  # 'docx' or 'pdf'
        
        if not address:
            return jsonify({'error': 'Address parameter is required'}), 400
        
        # If lat/lng not provided, geocode the address
        if lat is None or lng is None:
            geocode_url = 'https://maps.googleapis.com/maps/api/geocode/json'
            params = {
                'address': address,
                'key': Config.GOOGLE_MAPS_API_KEY,
                'components': 'country:US'
            }
            response = requests.get(geocode_url, params=params, timeout=10)
            data = response.json()
            
            if data['status'] != 'OK' or not data['results']:
                return jsonify({
                    'error': 'Could not geocode address',
                    'details': data.get('error_message', data.get('status', 'Unknown error'))
                }), 400
            
            location = data['results'][0]['geometry']['location']
            lat = location['lat']
            lng = location['lng']
            
            # Extract zip code if not provided
            if not zip_code:
                for component in data['results'][0]['address_components']:
                    if 'postal_code' in component['types']:
                        zip_code = component['long_name']
                        break
        
        db: Session = next(get_db())
        
        # Get census data for zip code
        census_record = None
        if zip_code:
            census_record = db.query(CensusData).options(load_only(*_CENSUS_LOAD_COLUMNS)).filter(CensusData.zip_code == zip_code).first()
        
        # STEP 1: Get "zoned" schools = nearest elementary, middle, high in school_data within ~5 miles (no Apify)
        zoned_schools = []
        zoned_school_names = set()
        search_radius_zoned = 5.0 / 69.0
        qp = {
            'lat': lat, 'lng': lng,
            'lat_min': lat - search_radius_zoned, 'lat_max': lat + search_radius_zoned,
            'lng_min': lng - search_radius_zoned, 'lng_max': lng + search_radius_zoned
        }
        dist_sql = "3959 * acos(cos(radians(:lat)) * cos(radians(latitude)) * cos(radians(longitude) - radians(:lng)) + sin(radians(:lat)) * sin(radians(latitude)))"
        q_elem = text("SELECT elementary_school_name, elementary_school_address, elementary_school_rating, latitude, longitude FROM school_data WHERE elementary_school_rating IS NOT NULL AND latitude BETWEEN :lat_min AND :lat_max AND longitude BETWEEN :lng_min AND :lng_max ORDER BY " + dist_sql + " LIMIT 1")
        q_mid = text("SELECT middle_school_name, middle_school_address, middle_school_rating, latitude, longitude FROM school_data WHERE middle_school_rating IS NOT NULL AND latitude BETWEEN :lat_min AND :lat_max AND longitude BETWEEN :lng_min AND :lng_max ORDER BY " + dist_sql + " LIMIT 1")
        q_high = text("SELECT high_school_name, high_school_address, high_school_rating, latitude, longitude FROM school_data WHERE high_school_rating IS NOT NULL AND latitude BETWEEN :lat_min AND :lat_max AND longitude BETWEEN :lng_min AND :lng_max ORDER BY " + dist_sql + " LIMIT 1")
        import math
        for query, name_col, addr_col, rating_col, school_type in [
            (q_elem, 0, 1, 2, 'Elementary'),
            (q_mid, 0, 1, 2, 'Middle'),
            (q_high, 0, 1, 2, 'High'),
        ]:
            row = db.execute(query, qp).fetchone()
            if row and row[name_col]:
                sch_lat, sch_lng = row[3], row[4]
                dist = 3959 * math.acos(
                    math.cos(math.radians(lat)) * math.cos(math.radians(sch_lat)) * math.cos(math.radians(sch_lng) - math.radians(lng)) +
                    math.sin(math.radians(lat)) * math.sin(math.radians(sch_lat))
                )
                zoned_schools.append({
                    'name': row[name_col],
                    'address': row[addr_col] or 'N/A',
                    'type': school_type,
                    'rating': row[rating_col],
                    'distance': dist,
                    'is_zoned': True
                })
                zoned_school_names.add(row[name_col].lower())

        # STEP 2: Get additional schools (sorted by rating, then distance) to fill up to 10 total
        # Mirror GreatSchools logic: 5-7 miles radius, public/charter only, rated schools only
        search_radius = 6.0 / 69.0  # ~6 miles radius (GreatSchools typically uses 5-7 miles)
        additional_query = text("""
            SELECT 
                elementary_school_name, elementary_school_address, elementary_school_rating,
                middle_school_name, middle_school_address, middle_school_rating,
                high_school_name, high_school_address, high_school_rating,
                latitude, longitude,
                3959 * acos(
                    cos(radians(:lat)) * cos(radians(latitude)) * 
                    cos(radians(longitude) - radians(:lng)) + 
                    sin(radians(:lat)) * sin(radians(latitude))
                ) as distance_miles
            FROM school_data
            WHERE (elementary_school_rating IS NOT NULL 
                   OR middle_school_rating IS NOT NULL 
                   OR high_school_rating IS NOT NULL)
              AND latitude BETWEEN :lat_min AND :lat_max
              AND longitude BETWEEN :lng_min AND :lng_max
        """)
        
        query_params = {
            'lat': lat,
            'lng': lng,
            'lat_min': lat - search_radius,
            'lat_max': lat + search_radius,
            'lng_min': lng - search_radius,
            'lng_max': lng + search_radius
        }
        
        all_results = db.execute(additional_query, query_params).fetchall()
        
        # Extract all schools, excluding zoned ones
        additional_schools = []
        for row in all_results:
            distance = row[11]
            
            # Elementary
            if row[0] and row[2] is not None and row[0].lower() not in zoned_school_names:
                additional_schools.append({
                    'name': row[0],
                    'address': row[1] or 'N/A',
                    'type': 'Elementary',
                    'rating': row[2],
                    'distance': distance,
                    'is_zoned': False
                })
            
            # Middle
            if row[3] and row[5] is not None and row[3].lower() not in zoned_school_names:
                additional_schools.append({
                    'name': row[3],
                    'address': row[4] or 'N/A',
                    'type': 'Middle',
                    'rating': row[5],
                    'distance': distance,
                    'is_zoned': False
                })
            
            # High
            if row[6] and row[8] is not None and row[6].lower() not in zoned_school_names:
                additional_schools.append({
                    'name': row[6],
                    'address': row[7] or 'N/A',
                    'type': 'High',
                    'rating': row[8],
                    'distance': distance,
                    'is_zoned': False
                })
        
        # Sort additional schools by rating (descending), then distance (ascending)
        # Mirror GreatSchools: Filter for public/charter schools only, exclude private and unrated
        # Note: We'll filter by school type if we have that data, but for now we'll rely on rating filter
        additional_schools = [s for s in additional_schools if s['rating'] is not None]  # Exclude unrated schools
        additional_schools.sort(key=lambda x: (-x['rating'] if x['rating'] is not None else 0, x['distance'] if x['distance'] is not None else float('inf')))
        
        # Combine: zoned schools first, then additional schools up to 10 total
        schools = zoned_schools + additional_schools[:10 - len(zoned_schools)]
        
        # Format for display
        formatted_schools = []
        for school in schools:
            formatted_schools.append({
                'name': school['name'],
                'address': school['address'],
                'type': school['type'],
                'rating': f"{school['rating']:.1f}/10" if school['rating'] is not None else 'N/A',
                'distance': f"{school['distance']:.2f} miles" if school['distance'] is not None else 'N/A',
                'is_zoned': school.get('is_zoned', False)
            })
        
        schools = formatted_schools
        
        # Generate document
        if format_type == 'pdf':
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib import colors
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            story = []
            styles = getSampleStyleSheet()
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.HexColor('#1a73e8'),
                spaceAfter=30
            )
            story.append(Paragraph("Site Selection Report", title_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Demographics Section
            story.append(Paragraph("<b>1. Demographics</b>", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))
            
            demo_data = [
                ['Field', 'Value'],
                ['Address', address],
                ['Zip Code', zip_code or 'N/A'],
                ['Population', f"{census_record.population:,}" if census_record and census_record.population else 'N/A'],
                ['Median Household Income (MHI)', f"${census_record.average_household_income:,.0f}" if census_record and census_record.average_household_income else 'N/A'],
                ['Median Age', f"{census_record.median_age:.1f} years" if census_record and census_record.median_age else 'N/A']
            ]
            
            demo_table = Table(demo_data, colWidths=[2.5*inch, 4*inch])
            demo_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(demo_table)
            story.append(Spacer(1, 0.3*inch))
            
            # Schools Section
            story.append(Paragraph("<b>2. Great School Scores (Top 10 Schools)</b>", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))
            
            if schools:
                # Create table with styled cells for zoned schools
                from reportlab.platypus import Table, TableStyle, Paragraph
                from reportlab.lib.styles import ParagraphStyle
                
                school_data = [['School Name', 'Address', 'Type', 'Rating', 'Proximity']]
                zoned_row_indices = []  # Track which rows are zoned
                
                for i, school in enumerate(schools):
                    row_idx = i + 1  # +1 for header row
                    school_data.append([
                        school['name'],
                        school['address'],
                        school['type'],
                        school['rating'],
                        school['distance']
                    ])
                    if school.get('is_zoned', False):
                        zoned_row_indices.append(row_idx)
                
                school_table = Table(school_data, colWidths=[1.8*inch, 1.8*inch, 0.8*inch, 0.8*inch, 0.8*inch])
                
                # Base table style
                table_style = [
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]
                
                # Style zoned schools with bold and blue color
                for row_idx in zoned_row_indices:
                    table_style.extend([
                        ('FONTNAME', (0, row_idx), (-1, row_idx), 'Helvetica-Bold'),
                        ('TEXTCOLOR', (0, row_idx), (-1, row_idx), colors.HexColor('#1a73e8')),
                        ('BACKGROUND', (0, row_idx), (-1, row_idx), colors.HexColor('#e8f0fe')),
                    ])
                
                # Add alternating row colors for non-zoned rows
                non_zoned_rows = [i for i in range(1, len(school_data)) if i not in zoned_row_indices]
                for i, row_idx in enumerate(non_zoned_rows):
                    bg_color = colors.white if i % 2 == 0 else colors.lightgrey
                    table_style.append(('BACKGROUND', (0, row_idx), (-1, row_idx), bg_color))
                
                school_table.setStyle(TableStyle(table_style))
                story.append(school_table)
            else:
                story.append(Paragraph("No school data available for this location.", styles['Normal']))
            
            doc.build(story)
            buffer.seek(0)
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'site_report_{zip_code or "unknown"}_{timestamp}.pdf'
            
            return Response(
                buffer.getvalue(),
                mimetype='application/pdf',
                headers={'Content-Disposition': f'attachment; filename="{filename}"'}
            )
        else:
            # Word document (docx)
            try:
                from docx import Document
                from docx.shared import Inches, Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                doc = Document()
                
                # Title
                title = doc.add_heading('Site Selection Report', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title.runs[0].font.color.rgb = RGBColor(26, 115, 232)
                
                # Demographics Section
                doc.add_heading('1. Demographics', 1)
                
                demo_data = [
                    ('Field', 'Value'),
                    ('Address', address),
                    ('Zip Code', zip_code or 'N/A'),
                    ('Population', f"{census_record.population:,}" if census_record and census_record.population else 'N/A'),
                    ('Median Household Income (MHI)', f"${census_record.average_household_income:,.0f}" if census_record and census_record.average_household_income else 'N/A'),
                    ('Median Age', f"{census_record.median_age:.1f} years" if census_record and census_record.median_age else 'N/A')
                ]
                
                demo_table = doc.add_table(rows=6, cols=2)
                demo_table.style = 'Light Grid Accent 1'
                
                for i, (field, value) in enumerate(demo_data):
                    row = demo_table.rows[i]
                    row.cells[0].text = field
                    row.cells[1].text = str(value)
                    if i == 0:  # Header row
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                
                doc.add_paragraph()  # Spacing
                
                # Schools Section
                doc.add_heading('2. Great School Scores (Top 10 Schools)', 1)
                
                if schools:
                    school_table = doc.add_table(rows=1, cols=5)
                    school_table.style = 'Light Grid Accent 1'
                    
                    # Header row
                    header_cells = school_table.rows[0].cells
                    headers = ['School Name', 'Address', 'Type', 'Rating', 'Proximity']
                    for i, header in enumerate(headers):
                        header_cells[i].text = header
                        for paragraph in header_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    # Data rows - style zoned schools with bold and blue color
                    for school in schools:
                        row_cells = school_table.add_row().cells
                        is_zoned = school.get('is_zoned', False)
                        
                        # Set cell values and style
                        for i, value in enumerate([school['name'], school['address'], school['type'], school['rating'], school['distance']]):
                            row_cells[i].text = value
                            
                            # Style zoned schools
                            if is_zoned:
                                for paragraph in row_cells[i].paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
                                        run.font.color.rgb = RGBColor(26, 115, 232)  # Blue color
                                    # Set cell background color (light blue)
                                    shading_elm = paragraph._element.get_or_add_pPr().get_or_add_shd()
                                    shading_elm.set('fill', 'E8F0FE')  # Light blue background
                else:
                    doc.add_paragraph('No school data available for this location.')
                
                # Save to BytesIO
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f'site_report_{zip_code or "unknown"}_{timestamp}.docx'
                
                return Response(
                    buffer.getvalue(),
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    headers={'Content-Disposition': f'attachment; filename="{filename}"'}
                )
            except ImportError:
                return jsonify({
                    'error': 'python-docx not installed',
                    'message': 'Install with: pip install python-docx'
                }), 500
        
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"ERROR in export_report: {e}")
        print(error_trace)
        return jsonify({
            'error': str(e),
            'message': 'Error generating report',
            'traceback': error_trace
        }), 500


@api.route('/schools/address', methods=['GET'])
def get_schools_by_address():
    """Get school ratings for an address. Uses nearest schools in school_data within ~5 miles (no Apify)."""
    try:
        import requests
        from config.config import Config
        from sqlalchemy import text

        address = request.args.get('address')
        lat = request.args.get('lat', type=float)
        lng = request.args.get('lng', type=float)

        if not address:
            return jsonify({'error': 'Address parameter is required'}), 400

        # Geocode if lat/lng not provided
        if lat is None or lng is None:
            geocode_url = 'https://maps.googleapis.com/maps/api/geocode/json'
            params = {
                'address': address,
                'key': Config.GOOGLE_MAPS_API_KEY,
                'components': 'country:US'
            }
            response = requests.get(geocode_url, params=params, timeout=10)
            data = response.json()
            if data['status'] != 'OK' or not data['results']:
                return jsonify({
                    'error': 'Could not geocode address',
                    'details': data.get('error_message', data.get('status', 'Unknown error'))
                }), 400
            location = data['results'][0]['geometry']['location']
            lat = location['lat']
            lng = location['lng']

        # Distance-based only: nearest schools in school_data within ~5 miles (no Apify)
        db: Session = next(get_db())
        search_radius = 5.0 / 69.0
        query_params = {
            'lat': lat,
            'lng': lng,
            'lat_min': lat - search_radius,
            'lat_max': lat + search_radius,
            'lng_min': lng - search_radius,
            'lng_max': lng + search_radius
        }

        elementary_query = text("""
            SELECT elementary_school_name, elementary_school_rating, elementary_school_address
            FROM school_data
            WHERE elementary_school_rating IS NOT NULL
              AND latitude BETWEEN :lat_min AND :lat_max
              AND longitude BETWEEN :lng_min AND :lng_max
            ORDER BY 3959 * acos(
                cos(radians(:lat)) * cos(radians(latitude)) * cos(radians(longitude) - radians(:lng)) +
                sin(radians(:lat)) * sin(radians(latitude))
            )
            LIMIT 1
        """)
        middle_query = text("""
            SELECT middle_school_name, middle_school_rating, middle_school_address
            FROM school_data
            WHERE middle_school_rating IS NOT NULL
              AND latitude BETWEEN :lat_min AND :lat_max
              AND longitude BETWEEN :lng_min AND :lng_max
            ORDER BY 3959 * acos(
                cos(radians(:lat)) * cos(radians(latitude)) * cos(radians(longitude) - radians(:lng)) +
                sin(radians(:lat)) * sin(radians(latitude))
            )
            LIMIT 1
        """)
        high_query = text("""
            SELECT high_school_name, high_school_rating, high_school_address
            FROM school_data
            WHERE high_school_rating IS NOT NULL
              AND latitude BETWEEN :lat_min AND :lat_max
              AND longitude BETWEEN :lng_min AND :lng_max
            ORDER BY 3959 * acos(
                cos(radians(:lat)) * cos(radians(latitude)) * cos(radians(longitude) - radians(:lng)) +
                sin(radians(:lat)) * sin(radians(latitude))
            )
            LIMIT 1
        """)

        elem_result = db.execute(elementary_query, query_params).fetchone()
        mid_result = db.execute(middle_query, query_params).fetchone()
        high_result = db.execute(high_query, query_params).fetchone()

        elementary_name = elem_result[0] if elem_result else None
        elementary_rating = elem_result[1] if elem_result else None
        elementary_addr = elem_result[2] if elem_result else None
        middle_name = mid_result[0] if mid_result else None
        middle_rating = mid_result[1] if mid_result else None
        middle_addr = mid_result[2] if mid_result else None
        high_name = high_result[0] if high_result else None
        high_rating = high_result[1] if high_result else None
        high_addr = high_result[2] if high_result else None

        # Blended score
        ratings = [r for r in [elementary_rating, middle_rating, high_rating] if r is not None]
        blended_score = sum(ratings) / len(ratings) if ratings else None

        # Zip from address
        zip_code = None
        if address:
            import re
            zip_match = re.search(r'\b\d{5}(?:-\d{4})?\b', address)
            if zip_match:
                zip_code = zip_match.group(0)

        result = {
            'zip_code': zip_code,
            'address': address,
            'latitude': lat,
            'longitude': lng,
            'elementary_school_name': elementary_name,
            'elementary_school_rating': elementary_rating,
            'elementary_school_address': elementary_addr,
            'middle_school_name': middle_name,
            'middle_school_rating': middle_rating,
            'middle_school_address': middle_addr,
            'high_school_name': high_name,
            'high_school_rating': high_rating,
            'high_school_address': high_addr,
            'blended_school_score': blended_score,
            'school_source': 'distance_fallback',
        }
        return jsonify(result)
        
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"ERROR in get_schools_by_address: {e}")
        print(error_trace)
        return jsonify({
            'error': str(e),
            'message': 'Error fetching school data',
            'traceback': error_trace
        }), 500


@api.route('/schools/address/all-zoned', methods=['GET'])
def get_all_zoned_schools_for_address():
    """
    Return ALL NCES attendance zones that contain the given address (point-in-polygon).
    Use for dropdown/export: list every school the address is zoned for (NC/SC only).
    """
    try:
        import requests
        from config.config import Config

        address = request.args.get('address')
        lat = request.args.get('lat', type=float)
        lng = request.args.get('lng', type=float)

        if lat is None or lng is None:
            if not address:
                return jsonify({'error': 'Provide address= or lat= and lng='}), 400
            geocode_url = 'https://maps.googleapis.com/maps/api/geocode/json'
            params = {'address': address, 'key': Config.GOOGLE_MAPS_API_KEY, 'components': 'country:US'}
            response = requests.get(geocode_url, params=params, timeout=10)
            data = response.json()
            if data.get('status') != 'OK' or not data.get('results'):
                return jsonify({'error': 'Could not geocode address'}), 400
            loc = data['results'][0]['geometry']['location']
            lat, lng = loc['lat'], loc['lng']

        db: Session = next(get_db())
        zones = db.query(AttendanceZone).filter(
            or_(AttendanceZone.state == 'NC', AttendanceZone.state == 'SC')
        ).all()
        if not zones:
            return jsonify({
                'address': address,
                'latitude': lat,
                'longitude': lng,
                'elementary': [], 'middle': [], 'high': [],
                'message': 'No NCES attendance zones loaded (NC/SC only).'
            })

        zones_list = [z.to_dict() for z in zones]
        by_level = find_all_zoned_schools(lat, lng, zones_list)

        def to_summary(zone_list):
            return [{'school_name': z.get('school_name'), 'school_level': z.get('school_level'),
                     'school_district': z.get('school_district'), 'state': z.get('state')} for z in zone_list]

        return jsonify({
            'address': address,
            'latitude': lat,
            'longitude': lng,
            'elementary': to_summary(by_level.get('elementary', [])),
            'middle': to_summary(by_level.get('middle', [])),
            'high': to_summary(by_level.get('high', [])),
        })
    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'traceback': traceback.format_exc()}), 500


def _rating_for_school(db: Session, school_name: str, level: str) -> Optional[float]:
    """Look up school rating from school_data by name and level. Returns rating 1-10 or None."""
    if not school_name or not level:
        return None
    level = level.lower()
    if level == 'elementary':
        row = db.query(SchoolData).filter(
            SchoolData.elementary_school_name.ilike(f'%{school_name}%'),
            SchoolData.elementary_school_rating.isnot(None)
        ).first()
        return float(row.elementary_school_rating) if row else None
    if level == 'middle':
        row = db.query(SchoolData).filter(
            SchoolData.middle_school_name.ilike(f'%{school_name}%'),
            SchoolData.middle_school_rating.isnot(None)
        ).first()
        return float(row.middle_school_rating) if row else None
    if level == 'high':
        row = db.query(SchoolData).filter(
            SchoolData.high_school_name.ilike(f'%{school_name}%'),
            SchoolData.high_school_rating.isnot(None)
        ).first()
        return float(row.high_school_rating) if row else None
    return None


def _school_info_for_name_level(
    db: Session, school_name: str, level: str
) -> Optional[Dict]:
    """
    Look up school_data row by name and level. Returns dict with name, rating, address
    (or None for each). Used so UI can show zoned school name even when rating is missing.
    """
    if not school_name or not level:
        return None
    level = level.lower()
    if level == 'elementary':
        row = db.query(SchoolData).filter(
            SchoolData.elementary_school_name.ilike(f'%{school_name}%')
        ).first()
        if not row:
            return None
        return {
            'name': row.elementary_school_name,
            'rating': float(row.elementary_school_rating) if row.elementary_school_rating is not None else None,
            'address': row.elementary_school_address,
        }
    if level == 'middle':
        row = db.query(SchoolData).filter(
            SchoolData.middle_school_name.ilike(f'%{school_name}%')
        ).first()
        if not row:
            return None
        return {
            'name': row.middle_school_name,
            'rating': float(row.middle_school_rating) if row.middle_school_rating is not None else None,
            'address': row.middle_school_address,
        }
    if level == 'high':
        row = db.query(SchoolData).filter(
            SchoolData.high_school_name.ilike(f'%{school_name}%')
        ).first()
        if not row:
            return None
        return {
            'name': row.high_school_name,
            'rating': float(row.high_school_rating) if row.high_school_rating is not None else None,
            'address': row.high_school_address,
        }
    return None


@api.route('/zips/<zip_code>/school-zones', methods=['GET'])
def get_school_zones_by_zip(zip_code: str):
    """
    For a zip code: list school districts that touch the zip, schools per district,
    district strength (avg rating), and GeoJSON geometry for each district's slice of the zip.
    NC/SC only (attendance zones). Requires zip boundary in data/zip_boundaries/{zip}.geojson.
    """
    try:
        zip_polygon = load_zip_polygon(zip_code)
        if zip_polygon is None:
            return jsonify({
                'error': 'Zip boundary not found',
                'message': f'No boundary for zip {zip_code}. Run: python scripts/download_accurate_boundaries.py --zip-codes {zip_code}'
            }), 404

        db: Session = next(get_db())
        zones = db.query(AttendanceZone).filter(
            or_(AttendanceZone.state == 'NC', AttendanceZone.state == 'SC')
        ).all()
        if not zones:
            return jsonify({
                'zip_code': zip_code,
                'district_count': 0,
                'districts': [],
                'message': 'No NCES attendance zones loaded (NC/SC only).'
            })

        zones_list = [z.to_dict() for z in zones]
        intersecting, diag = zones_intersecting_zip_diagnostic(zip_polygon, zones_list)
        if not intersecting:
            return jsonify({
                'zip_code': zip_code,
                'district_count': 0,
                'districts': [],
                'message': 'No attendance zones intersect this zip (NC/SC data only).',
                'debug': diag,
            })

        grouped = group_zones_by_district(intersecting)
        DISTRICT_COLORS = ['#4A90D9', '#50C878', '#E6A23C', '#E07070', '#9B59B6', '#1ABC9C', '#E67E22', '#3498DB']

        districts_out = []
        for i, grp in enumerate(grouped):
            district_id = grp['district_id']
            district_zones = grp['zones']
            schools = []
            ratings = []
            for z in district_zones:
                name = z.get('school_name') or 'Unknown'
                level = (z.get('school_level') or 'unknown').lower()
                rating = _rating_for_school(db, name, level)
                schools.append({'name': name, 'level': level, 'rating': rating})
                if rating is not None:
                    ratings.append(rating)
            avg_rating = sum(ratings) / len(ratings) if ratings else None
            geometry = district_geometry_in_zip(zip_polygon, district_zones)
            color = DISTRICT_COLORS[i % len(DISTRICT_COLORS)]
            districts_out.append({
                'district_id': district_id,
                'district_name': grp['district_name'],
                'schools': schools,
                'avg_rating': round(avg_rating, 1) if avg_rating is not None else None,
                'geometry': geometry,
                'color': color,
            })

        districts_out.sort(key=lambda d: (d['avg_rating'] is None, -(d['avg_rating'] or 0)))

        return jsonify({
            'zip_code': zip_code,
            'district_count': len(districts_out),
            'districts': districts_out,
        })
    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'traceback': traceback.format_exc()}), 500


@api.route('/schools/zip/<zip_code>', methods=['GET'])
def get_schools_by_zip(zip_code: str):
    """Get school ratings for a zip code."""
    try:
        db: Session = next(get_db())
        
        # Check if we have cached data
        cached = db.query(SchoolData).filter(SchoolData.zip_code == zip_code).first()
        if cached:
            return jsonify(cached.to_dict())
        
        # If not cached, we need an address to geocode
        # For now, return error suggesting to use address endpoint
        return jsonify({
            'error': 'School data not found for this zip code',
            'message': 'Please use /api/schools/address?address=<full_address> to fetch school data'
        }), 404
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500
